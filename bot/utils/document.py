import aiohttp
import io
import docx
from docx.shared import Inches
from bot.utils.logger import get_logger
from config.config import load_config
from cachetools import TTLCache
import asyncio
import random
import json

logger = get_logger(__name__)
config = load_config()
question_cache = TTLCache(maxsize=1000, ttl=3600)

async def generate_test_questions(subject, description, questions_count):
    cache_key = f"{subject}:{description}:{questions_count}"
    cached_questions = question_cache.get(cache_key)
    if cached_questions:
        return cached_questions
    fallback_questions = {
        "Matematika": [
            {"question": "2 + 2 = ?", "options": ["3", "4", "5", "6"], "answer": "4"},
            {"question": "x^2 - 4 = 0 tenglamaning yechimi nima?", "options": ["x = ±2", "x = ±4", "x = 0", "x = ±1"], "answer": "x = ±2"}
        ],
        "Tarix": [
            {"question": "Ikkinchi jahon urushi qachon boshlangan?", "options": ["1939", "1941", "1945", "1935"], "answer": "1939"},
            {"question": "Amir Temur qachon tug'ilgan?", "options": ["1336", "1340", "1320", "1350"], "answer": "1336"}
        ],
        "Biologiya": [
            {"question": "DNK molekulasining asosiy vazifasi nima?", "options": ["Energiya saqlash", "Ma'lumot saqlash", "Transport", "Struktura"], "answer": "Ma'lumot saqlash"},
            {"question": "Fotosintez jarayoni qayerda sodir bo'ladi?", "options": ["Mitoxondriya", "Xloroplast", "Yadro", "Ribosoma"], "answer": "Xloroplast"}
        ]
    }
    try:
        prompt = (
            f"Create {questions_count} multiple-choice test questions in Uzbek for the subject '{subject}'"
            f"{f' with the topic: {description}' if description else ''}. "
            "Each question should have 4 options and one correct answer. "
            "Return the result as a JSON list where each item has 'question', 'options' (list), and 'answer' fields."
        )
        headers = {
            "Authorization": f"Bearer {config.openrouter_api_key}",
            "Content-Type": "application/json"
        }
        payload = {
            "model": "deepseek/deepseek-prover-v2:free",
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": 4096,
            "temperature": 0.7
        }
        async with aiohttp.ClientSession() as session:
            async with session.post(
                "https://openrouter.ai/api/v1/chat/completions",
                json=payload,
                headers=headers
            ) as response:
                if response.status != 200:
                    logger.error(f"OpenRouter API error: {response.status} - {await response.text()}")
                    subject_key = subject.capitalize()
                    available_questions = fallback_questions.get(subject_key, [])
                    return random.sample(available_questions, min(questions_count, len(available_questions)))
                result = await response.json()
                questions_text = result["choices"][0]["message"]["content"]
                questions = json.loads(questions_text)
                if len(questions) < questions_count:
                    subject_key = subject.capitalize()
                    available_questions = fallback_questions.get(subject_key, [])
                    questions.extend(random.sample(available_questions, min(questions_count - len(questions), len(available_questions))))
                question_cache[cache_key] = questions
                return questions
    except Exception as e:
        logger.error(f"Error generating questions: {e}", exc_info=True)
        subject_key = subject.capitalize()
        available_questions = fallback_questions.get(subject_key, [])
        return random.sample(available_questions, min(questions_count, len(available_questions)))

async def generate_test_document(subject, description, questions_count):
    questions = await generate_test_questions(subject, description, questions_count)
    doc = docx.Document()
    doc.add_heading(f"{subject} bo'yicha test", 0)
    if description:
        doc.add_paragraph(f"Tavsif: {description}")
    for i, q in enumerate(questions, 1):
        doc.add_paragraph(f"{i}. {q['question']}")
        for j, option in enumerate(q['options'], 1):
            doc.add_paragraph(f"   {chr(96 + j)}) {option}")
        doc.add_paragraph(f"Javob: {q['answer']}\n")
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
